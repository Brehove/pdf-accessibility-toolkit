#!/usr/bin/env python3
"""
Convert Markdown to Accessible DOCX

Creates a fully accessible Word document with:
- Proper heading styles (H1, H2, H3)
- Document structure tags for screen readers
- Alt text on images
- Bulleted and numbered lists
- Document metadata

Usage:
    python md_to_accessible_docx.py input.md
    python md_to_accessible_docx.py input.md -o output.docx
    python md_to_accessible_docx.py *.md  # Batch process
"""

import argparse
import math
import re
import os
import base64
from typing import Optional
from dotenv import load_dotenv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# WCAG-compliant heading color (black)
HEADING_COLOR = RGBColor(0x00, 0x00, 0x00)

# Unicode superscript mapping for LaTeX cleanup
_SUPERSCRIPT_MAP = str.maketrans('0123456789,*', '⁰¹²³⁴⁵⁶⁷⁸⁹˙*')


def load_env_context(env_file: Optional[str] = None, input_files: Optional[list[str]] = None):
    """Load MISTRAL_API_KEY from likely .env locations."""
    script_dir = Path(__file__).resolve().parent
    candidates = []

    if env_file:
        candidates.append(Path(env_file).expanduser())

    if input_files:
        for file_path in input_files:
            parent = Path(file_path).expanduser().parent
            candidates.append(parent / ".env")

    candidates.extend([
        Path.cwd() / ".env",
        script_dir / ".env",
        script_dir.parent / ".env",
    ])

    loaded = []
    seen = set()
    for candidate in candidates:
        candidate = candidate.expanduser()
        key = str(candidate)
        if key in seen:
            continue
        seen.add(key)
        if candidate.is_file():
            load_dotenv(candidate, override=False)
            loaded.append(candidate)

    if loaded:
        print(f"Loaded .env: {loaded[0]}")
    else:
        print("No .env file found in default locations; using shell environment only.")


def clean_latex_notation(text):
    """Clean LaTeX math notation from OCR output.

    Handles common patterns from academic paper OCR:
    - $①$ (ORCID icons) → removed
    - $^{1,2,3}$ (superscript affiliations) → Unicode superscripts
    - $_{text}$ (subscripts) → plain text
    - $...$ (other inline math) → content without dollar signs
    """
    # Remove ORCID circled numbers: $①$ , $②$ etc.
    text = re.sub(r'\$[①②③④⑤⑥⑦⑧⑨⑩]+\$\s*,?\s*', '', text)

    # Convert superscript affiliations: $^{1,2,3}$ → ¹˙²˙³
    def _super(m):
        content = m.group(1)
        return content.translate(_SUPERSCRIPT_MAP)
    text = re.sub(r'\$\^{([0-9,*]+)}\$', _super, text)

    # Convert general superscripts: $^{text}$ → text
    text = re.sub(r'\$\^{(.*?)}\$', r'\1', text)

    # Convert subscripts: $_{text}$ → text
    text = re.sub(r'\$_{(.*?)}\$', r'\1', text)

    # Strip remaining inline math: $text$ → text
    text = re.sub(r'\$([^$]+)\$', r'\1', text)

    # Clean up extra whitespace left behind
    text = re.sub(r'  +', ' ', text)

    return text


def fix_table_headers(tables):
    """
    Fix table headers for accessibility.
    Sets first row as header row with proper markup.
    Returns number of tables fixed.
    """
    tables_fixed = 0
    for table in tables:
        # Set tblHeader on the first row
        header_row = table.rows[0]
        tr = header_row._tr
        trPr = tr.get_or_add_trPr()

        # Remove existing tblHeader elements
        for elem in list(trPr):
            if elem.tag == qn('w:tblHeader'):
                trPr.remove(elem)

        # Insert tblHeader
        tblHeader = OxmlElement('w:tblHeader')
        trPr.insert(0, tblHeader)

        # Set table-level properties
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblLook = tblPr.find(qn('w:tblLook'))
        if tblLook is None:
            tblLook = OxmlElement('w:tblLook')
            tblPr.append(tblLook)

        # Set accessibility attributes
        tblLook.set(qn('w:firstRow'), '1')
        tblLook.set(qn('w:lastRow'), '0')
        tblLook.set(qn('w:firstColumn'), '0')
        tblLook.set(qn('w:lastColumn'), '0')
        tblLook.set(qn('w:noHBand'), '0')
        tblLook.set(qn('w:noVBand'), '1')

        tables_fixed += 1

    return tables_fixed


def get_alt_text(image_path, alt_text_map=None):
    """Generate alt text for an image."""
    filename = Path(image_path).stem

    # Check custom alt text map if provided
    if alt_text_map:
        match = re.search(r'img_(\d+)', filename)
        if match:
            img_key = f"img_{match.group(1)}"
            if img_key in alt_text_map:
                return alt_text_map[img_key]

    # Generate from filename
    clean_name = filename.replace('_', ' ').replace('-', ' ')
    return f"Image: {clean_name}"


def is_placeholder_alt_text(alt_text, image_ref):
    if not alt_text:
        return True
    alt_clean = alt_text.strip().lower()
    if not alt_clean:
        return True
    image_name = Path(image_ref).name.lower() if image_ref else ""
    image_stem = Path(image_ref).stem.lower() if image_ref else ""
    if alt_clean in {image_name, image_stem}:
        return True
    if alt_clean.endswith(('.png', '.jpg', '.jpeg', '.gif')):
        return True
    if re.match(r'^(img|image|figure|photo|screenshot)[-_ ]?\d*$', alt_clean):
        return True
    return False


def caption_to_alt_text(text):
    """Convert a figure caption line into concise alt text."""
    cleaned = re.sub(r'\s+', ' ', (text or "")).strip()
    if not cleaned:
        return None
    cleaned = re.sub(r'^(figure|fig\.?)\s*\d*[:.\-]\s*', '', cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.strip()
    if not cleaned:
        return None
    return cleaned


def get_following_figure_caption(elements, index):
    """Return next paragraph if it looks like a figure caption."""
    if index + 1 >= len(elements):
        return None
    next_type, next_content = elements[index + 1]
    if next_type != "paragraph":
        return None
    line = (next_content or "").strip()
    if re.match(r'^(figure|fig\.?)\s*\d*[:.\-]\s+', line, flags=re.IGNORECASE):
        return line
    return None


def get_mistral_client():
    api_key = os.environ.get("MISTRAL_API_KEY")
    if not api_key:
        return None
    try:
        from mistralai import Mistral
        return Mistral(api_key=api_key)
    except Exception:
        return None


def generate_alt_text_mistral(image_path, model):
    client = get_mistral_client()
    if not client:
        return None
    try:
        with open(image_path, "rb") as f:
            image_bytes = f.read()
        ext = image_path.suffix.lower().lstrip(".") or "png"
        data_url = f"data:image/{ext};base64,{base64.b64encode(image_bytes).decode('utf-8')}"
        prompt = (
            "Write concise, descriptive accessibility alt text for this image. "
            "Use 1-2 detailed sentences. No quotes."
        )
        response = client.chat.complete(
            model=model,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": data_url},
                ],
            }],
        )
        content = response.choices[0].message.content
        if isinstance(content, list):
            content = " ".join(part.get("text", "") for part in content if isinstance(part, dict))
        alt = (content or "").strip().replace("\n", " ")
        alt_final = alt or None
        return alt_final
    except Exception:
        return None


def set_document_properties(doc, md_file):
    """Set document properties based on filename."""
    title = md_file.stem.replace('_', ' ').replace('-', ' ').title()
    doc.core_properties.title = title
    doc.core_properties.subject = "Accessible Document"


def add_image_with_alt_text(doc, paragraph, image_path, alt_text, width=Inches(5.5)):
    """Add an image with proper alt text for accessibility."""
    if not os.path.exists(image_path):
        return False

    if alt_text:
        alt_text = re.sub(r'\s+', ' ', alt_text).strip()

    run = paragraph.add_run()
    picture = run.add_picture(str(image_path), width=width)

    # Add alt text to the image
    inline = picture._inline
    doc_pr = inline.docPr
    doc_pr.set('descr', alt_text or "")
    doc_pr.set('title', "")

    return True


def find_image_path(img_ref, md_file):
    """Find the actual image file path relative to the markdown file."""
    md_dir = md_file.parent

    # Handle relative paths in markdown
    if img_ref.startswith('./'):
        img_ref = img_ref[2:]

    # Try the path as-is first
    full_path = md_dir / img_ref
    if full_path.exists():
        return full_path

    # Try with different extensions
    base_path = full_path.with_suffix('')
    for ext in ['.jpg', '.jpeg', '.png', '.gif']:
        test_path = base_path.with_suffix(ext)
        if test_path.exists():
            return test_path

    # Try in extracted_images folder
    if 'extracted_images' in img_ref:
        base_name = Path(img_ref).stem
        images_folder = md_dir / "extracted_images"
        for ext in ['.jpg', '.jpeg', '.png', '.gif']:
            test_path = images_folder / (base_name + ext)
            if test_path.exists():
                return test_path

    return None


def split_into_pages(md_content):
    """Split markdown content into pages based on <!-- Page X --> markers."""
    page_pattern = r'<!-- Page \d+ -->'
    parts = re.split(page_pattern, md_content)

    pages = []
    for part in parts:
        part = part.strip()
        if part and part != '---':
            pages.append(part)

    return pages


def is_table_separator(line):
    """Check if a line is a markdown table separator."""
    if '|' not in line:
        return False
    stripped = line.strip()
    return bool(re.match(r'^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$', stripped))


def split_table_row(line):
    """Split a markdown table row into cells."""
    row = line.strip()
    if row.startswith('|'):
        row = row[1:]
    if row.endswith('|'):
        row = row[:-1]
    cells = [cell.strip().replace('&amp;', '&') for cell in row.split('|')]
    return cells


def parse_table(lines, start_index):
    """Parse a markdown table starting at start_index."""
    rows = []
    header_line = lines[start_index]
    rows.append(split_table_row(header_line))

    i = start_index + 2  # Skip header and separator
    while i < len(lines):
        raw_line = lines[i]
        stripped = raw_line.strip()

        if not stripped:
            break
        if stripped.startswith('#') or stripped == '---' or stripped.startswith('<!--'):
            break
        if is_table_separator(stripped):
            i += 1
            continue

        if '|' in raw_line and stripped.startswith('|'):
            rows.append(split_table_row(raw_line))
        else:
            # Continuation of the previous row's last cell
            if rows and rows[-1]:
                rows[-1][-1] = (rows[-1][-1].rstrip() + '\n' + stripped).strip()
            else:
                break

        i += 1

    return rows, i


def parse_content(content):
    """Parse markdown content into elements."""
    elements = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip horizontal rules
        if line == '---':
            i += 1
            continue

        # Skip standalone page numbers
        if re.match(r'^\d+$', line):
            i += 1
            continue

        # Table (must come before headings and lists)
        if line.startswith('|') and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_rows, i = parse_table(lines, i)
            if table_rows:
                elements.append(('table', table_rows))
            continue

        # Heading 1
        if line.startswith('# '):
            elements.append(('h1', line[2:].replace('&amp;', '&')))
            i += 1
            continue

        # Heading 2
        if line.startswith('## '):
            elements.append(('h2', line[3:].replace('&amp;', '&')))
            i += 1
            continue

        # Heading 3
        if line.startswith('### '):
            elements.append(('h3', line[4:].replace('&amp;', '&')))
            i += 1
            continue

        # Image
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            elements.append(('image', {
                "src": img_match.group(2),
                "alt": img_match.group(1)
            }))
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', line):
            list_items = []
            while i < len(lines):
                list_match = re.match(r'^\d+\.\s(.+)', lines[i].strip())
                if list_match:
                    list_items.append(list_match.group(1).replace('&amp;', '&'))
                    i += 1
                elif lines[i].strip() == '':
                    i += 1
                    break
                else:
                    break
            if list_items:
                elements.append(('numbered_list', list_items))
            continue

        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            list_items = []
            while i < len(lines):
                stripped = lines[i].strip()
                if stripped.startswith('- ') or stripped.startswith('* '):
                    list_items.append(stripped[2:].replace('&amp;', '&'))
                    i += 1
                elif stripped == '':
                    i += 1
                    break
                else:
                    break
            if list_items:
                elements.append(('bullet_list', list_items))
            continue

        # Regular paragraph
        elements.append(('paragraph', line.replace('&amp;', '&')))
        i += 1

    return elements


def _looks_like_author_name(text):
    """Heuristic check for author name lines."""
    if '@' in text:
        return False
    stripped = text.strip()
    if not stripped or len(stripped) > 80:
        return False
    tokens = re.findall(r"[A-Za-z][A-Za-z'`\-\.]*", stripped)
    if len(tokens) < 2 or len(tokens) > 7:
        return False
    return all(t[0].isupper() for t in tokens if t and t[0].isalpha())


def _looks_like_email(text):
    """Heuristic check for email lines."""
    return bool(re.search(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', text))


def normalize_first_page_author_block(elements):
    """
    Detect a first-page stacked author block and convert it to a side-by-side grid.
    Expected pattern: [Name, Email, Affiliation, Location] repeated.
    """
    if not elements:
        return elements

    title_idx = next((i for i, (t, _) in enumerate(elements) if t == 'h1'), None)
    if title_idx is None:
        return elements

    idx = title_idx + 1
    paragraph_block = []
    while idx < len(elements) and elements[idx][0] == 'paragraph':
        paragraph_block.append(elements[idx][1].strip())
        idx += 1

    if len(paragraph_block) < 8:
        return elements

    authors = []
    consumed = 0
    while consumed + 3 < len(paragraph_block):
        name = paragraph_block[consumed]
        email = paragraph_block[consumed + 1]
        affiliation = paragraph_block[consumed + 2]
        location = paragraph_block[consumed + 3]
        if (
            _looks_like_author_name(name)
            and _looks_like_email(email)
            and not _looks_like_email(affiliation)
            and not _looks_like_email(location)
        ):
            authors.append({
                "name": name,
                "email": email,
                "affiliation": affiliation,
                "location": location,
            })
            consumed += 4
            continue
        break

    # Only apply if this clearly looks like an author block.
    if len(authors) < 2:
        return elements
    if consumed < len(paragraph_block) - 1:
        return elements

    rewritten = elements[:title_idx + 1]
    rewritten.append(('author_grid', authors))
    for leftover in paragraph_block[consumed:]:
        rewritten.append(('paragraph', leftover))
    rewritten.extend(elements[idx:])
    return rewritten


def add_paragraph_with_formatting(doc, text):
    """Add a paragraph with bold/italic text handling."""
    para = doc.add_paragraph()

    # Handle bold (**text**) and italic (*text*)
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)
    return para


def create_accessible_docx(
    md_file,
    output_file=None,
    auto_alt=False,
    alt_model="pixtral-12b",
    preserve_page_breaks=False,
    use_author_grid=True,
):
    """Create an accessible Word document from markdown."""
    md_file = Path(md_file)

    if output_file is None:
        output_file = md_file.with_suffix('.docx')
    else:
        output_file = Path(output_file)

    print(f"Reading: {md_file}")

    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Clean LaTeX notation from OCR output
    md_content = clean_latex_notation(md_content)

    # Check if content has page markers (from OCR)
    has_pages = bool(re.search(r'<!-- Page \d+ -->', md_content))

    if has_pages:
        pages = split_into_pages(md_content)
        print(f"Found {len(pages)} pages")
    else:
        pages = [md_content]

    # Create document
    doc = Document()
    set_document_properties(doc, md_file)

    image_count = 0
    page_count = 0
    data_tables = []

    for page_num, page_content in enumerate(pages, 1):
        elements = parse_content(page_content)
        if use_author_grid and page_num == 1:
            elements = normalize_first_page_author_block(elements)

        if not elements:
            continue

        page_count += 1
        if has_pages:
            print(f"  Processing page {page_num}...")

        # Process each element
        skip_indices = set()
        for elem_index, (elem_type, content) in enumerate(elements):
            if elem_index in skip_indices:
                continue
            if elem_type == 'h1':
                heading = doc.add_heading(content, level=1)
                for run in heading.runs:
                    run.font.color.rgb = HEADING_COLOR

            elif elem_type == 'h2':
                heading = doc.add_heading(content, level=2)
                for run in heading.runs:
                    run.font.color.rgb = HEADING_COLOR

            elif elem_type == 'h3':
                heading = doc.add_heading(content, level=3)
                for run in heading.runs:
                    run.font.color.rgb = HEADING_COLOR

            elif elem_type == 'paragraph':
                add_paragraph_with_formatting(doc, content)

            elif elem_type == 'table':
                table_rows = content
                if not table_rows:
                    continue
                num_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=1, cols=num_cols)
                data_tables.append(table)
                first_row_cells = table.rows[0].cells
                for ci in range(num_cols):
                    text = table_rows[0][ci] if ci < len(table_rows[0]) else ''
                    first_row_cells[ci].text = text

                for row in table_rows[1:]:
                    row_cells = table.add_row().cells
                    for ci in range(num_cols):
                        text = row[ci] if ci < len(row) else ''
                        row_cells[ci].text = text

            elif elem_type == 'author_grid':
                authors = content
                if not authors:
                    continue
                num_cols = min(3, len(authors))
                num_rows = math.ceil(len(authors) / num_cols)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.autofit = True

                for idx, author in enumerate(authors):
                    row_i = idx // num_cols
                    col_i = idx % num_cols
                    cell = table.cell(row_i, col_i)
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    name_run = cell_para.add_run(author["name"])
                    name_run.bold = True

                    for field in ("email", "affiliation", "location"):
                        cell_para.add_run("\n" + author[field])

                for idx in range(len(authors), num_rows * num_cols):
                    row_i = idx // num_cols
                    col_i = idx % num_cols
                    table.cell(row_i, col_i).text = ''

            elif elem_type == 'numbered_list':
                for item in content:
                    doc.add_paragraph(item, style='List Number')

            elif elem_type == 'bullet_list':
                for item in content:
                    doc.add_paragraph(item, style='List Bullet')

            elif elem_type == 'image':
                if isinstance(content, dict):
                    img_ref = content.get("src")
                    alt_from_md = (content.get("alt") or "").strip()
                else:
                    img_ref = content
                    alt_from_md = ""

                img_path = find_image_path(img_ref, md_file) if img_ref else None
                if img_path:
                    image_count += 1
                    is_placeholder = is_placeholder_alt_text(alt_from_md, img_ref)
                    alt_text = alt_from_md if alt_from_md and not is_placeholder else None
                    should_generate = auto_alt and is_placeholder
                    if should_generate:
                        generated_alt = generate_alt_text_mistral(img_path, alt_model)
                        if generated_alt:
                            alt_text = generated_alt
                    if not alt_text:
                        caption = get_following_figure_caption(elements, elem_index)
                        alt_text = caption_to_alt_text(caption)
                        if alt_text:
                            skip_indices.add(elem_index + 1)
                    if not alt_text:
                        alt_text = get_alt_text(img_ref)

                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if add_image_with_alt_text(doc, para, img_path, alt_text):
                        print(f"    Added image: {img_path.name}")

        # Optional: add page break between OCR pages.
        if preserve_page_breaks and has_pages and page_num < len(pages):
            doc.add_page_break()

    # Fix table headers for accessibility
    tables_fixed = fix_table_headers(data_tables)

    # Save document
    doc.save(output_file)
    print(f"\n✓ Created: {output_file}")
    print(f"  Pages: {page_count}")
    if image_count:
        print(f"  Images: {image_count}")
    if tables_fixed:
        print(f"  Tables: {tables_fixed} (headers fixed)")


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to accessible DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python md_to_accessible_docx.py document.md
  python md_to_accessible_docx.py document.md -o accessible.docx
  python md_to_accessible_docx.py *.md
        """
    )
    parser.add_argument('files', nargs='+', help='Markdown file(s) to convert')
    parser.add_argument('-o', '--output', help='Output filename (single file only)')
    parser.add_argument('--env-file', help='Optional path to .env containing MISTRAL_API_KEY')
    parser.add_argument('--no-auto-alt', action='store_true',
                        help='Disable automatic alt text generation via Mistral API')
    parser.add_argument('--alt-model', default='pixtral-12b',
                        help='Mistral model for alt text generation (default: pixtral-12b)')
    parser.add_argument(
        '--preserve-page-breaks',
        action='store_true',
        help='Insert DOCX page breaks at OCR page markers (off by default)'
    )
    parser.add_argument(
        '--no-author-grid',
        action='store_true',
        help='Disable first-page author block side-by-side layout heuristic'
    )

    args = parser.parse_args()

    if args.output and len(args.files) > 1:
        print("Error: -o/--output can only be used with a single input file")
        return

    load_env_context(args.env_file, args.files)

    # Skip common documentation files
    skip_files = {'readme.md', 'changelog.md', 'contributing.md', 'license.md'}

    for md_file in args.files:
        md_path = Path(md_file)
        if md_path.name.lower() in skip_files:
            print(f"Skipping documentation file: {md_file}")
            continue
        if not md_path.exists():
            print(f"Error: File not found: {md_file}")
            continue
        if md_path.suffix.lower() != '.md':
            print(f"Skipping non-markdown file: {md_file}")
            continue

        output = args.output if args.output else None
        auto_alt = not args.no_auto_alt
        create_accessible_docx(
            md_path,
            output,
            auto_alt=auto_alt,
            alt_model=args.alt_model,
            preserve_page_breaks=args.preserve_page_breaks,
            use_author_grid=not args.no_author_grid,
        )
        print()


if __name__ == "__main__":
    main()
