#!/usr/bin/env python3
"""
Fix DOCX Table Headers for Accessibility
Ensures all tables have properly marked header rows for WCAG 2.1 AA compliance.

Usage:
    python fix_docx_table_headers.py input.docx
    python fix_docx_table_headers.py input.docx -o output.docx
    python fix_docx_table_headers.py *.docx  # Batch process multiple files
"""

import argparse
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_table_header_row(table, header_row_index=0):
    """
    Set a row as the header row with all necessary accessibility properties.

    Args:
        table: python-docx Table object
        header_row_index: Which row is the header (default: 0 = first row)

    Returns:
        List of header cell texts for verification
    """
    # 1. Set tblHeader on the header row
    header_row = table.rows[header_row_index]
    tr = header_row._tr
    trPr = tr.get_or_add_trPr()

    # Remove existing tblHeader elements
    for elem in list(trPr):
        if elem.tag == qn('w:tblHeader'):
            trPr.remove(elem)

    # Insert tblHeader at beginning (order matters for some parsers)
    tblHeader = OxmlElement('w:tblHeader')
    trPr.insert(0, tblHeader)

    # 2. Set table-level properties (tblLook)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Get or create tblLook
    tblLook = tblPr.find(qn('w:tblLook'))
    if tblLook is None:
        tblLook = OxmlElement('w:tblLook')
        tblPr.append(tblLook)

    # Set all tblLook attributes for accessibility
    tblLook.set(qn('w:firstRow'), '1')      # First row is header (critical!)
    tblLook.set(qn('w:lastRow'), '0')
    tblLook.set(qn('w:firstColumn'), '0')
    tblLook.set(qn('w:lastColumn'), '0')
    tblLook.set(qn('w:noHBand'), '0')
    tblLook.set(qn('w:noVBand'), '1')

    # Return header cell texts for verification
    header_cells = [cell.text[:40].strip() for cell in header_row.cells]
    # Remove duplicates (merged cells)
    unique_headers = list(dict.fromkeys(header_cells))

    return unique_headers


def fix_docx_tables(input_path, output_path=None):
    """
    Fix all table headers in a DOCX file.

    Args:
        input_path: Path to input DOCX file
        output_path: Path to output file (default: input_accessible.docx)

    Returns:
        Tuple of (output_path, num_tables_fixed)
    """
    input_path = Path(input_path)

    if output_path is None:
        output_path = input_path.parent / f"{input_path.stem}_accessible{input_path.suffix}"
    else:
        output_path = Path(output_path)

    doc = Document(input_path)

    tables_fixed = 0
    for i, table in enumerate(doc.tables):
        headers = set_table_header_row(table)
        tables_fixed += 1
        print(f"  Table {i+1}: {headers}")

    doc.save(output_path)

    return output_path, tables_fixed


def verify_table_headers(docx_path):
    """Verify that table headers are properly set."""
    doc = Document(docx_path)

    results = []
    for i, table in enumerate(doc.tables):
        tr = table.rows[0]._tr
        trPr = tr.find(qn('w:trPr'))
        has_tblHeader = trPr is not None and trPr.find(qn('w:tblHeader')) is not None

        tblPr = table._tbl.find(qn('w:tblPr'))
        tblLook = tblPr.find(qn('w:tblLook')) if tblPr is not None else None
        firstRow = tblLook.get(qn('w:firstRow')) if tblLook is not None else None

        results.append({
            'table': i + 1,
            'tblHeader': has_tblHeader,
            'firstRow': firstRow == '1'
        })

    return results


def main():
    parser = argparse.ArgumentParser(
        description='Fix DOCX table headers for accessibility (WCAG 2.1 AA)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python fix_docx_table_headers.py document.docx
  python fix_docx_table_headers.py document.docx -o fixed.docx
  python fix_docx_table_headers.py *.docx
        """
    )
    parser.add_argument('files', nargs='+', help='DOCX file(s) to process')
    parser.add_argument('-o', '--output', help='Output file (only for single file input)')
    parser.add_argument('-v', '--verify', action='store_true', help='Verify headers after fixing')

    args = parser.parse_args()

    if args.output and len(args.files) > 1:
        print("Error: -o/--output can only be used with a single input file")
        return 1

    print("=" * 60)
    print("DOCX Table Header Fixer")
    print("=" * 60)

    total_tables = 0

    for file_path in args.files:
        file_path = Path(file_path)

        if not file_path.exists():
            print(f"\nSkipping {file_path} (not found)")
            continue

        if not file_path.suffix.lower() == '.docx':
            print(f"\nSkipping {file_path} (not a .docx file)")
            continue

        print(f"\nProcessing: {file_path.name}")

        try:
            output_path, num_tables = fix_docx_tables(
                file_path,
                args.output if args.output else None
            )
            total_tables += num_tables
            print(f"  ✓ Fixed {num_tables} table(s)")
            print(f"  ✓ Saved: {output_path.name}")

            if args.verify:
                print("  Verifying...")
                results = verify_table_headers(output_path)
                for r in results:
                    status = "✓" if r['tblHeader'] and r['firstRow'] else "✗"
                    print(f"    Table {r['table']}: {status} tblHeader={r['tblHeader']}, firstRow={r['firstRow']}")

        except Exception as e:
            print(f"  ✗ Error: {e}")

    print("\n" + "=" * 60)
    print(f"Complete! Fixed {total_tables} table(s) in {len(args.files)} file(s)")
    print("=" * 60)

    return 0


if __name__ == "__main__":
    exit(main())
